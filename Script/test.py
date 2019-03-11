# -*- coding: utf-8 -*-
import docx
import os
import GoogleTranslateUtils
import ExcelUtils
import WordUtils



class test:

    def __init__(self):
        pass

if __name__ == "__main__":



    # 当前的思路如下
    # 1. 在网上找到当天发布的一些内容（母语内容，如英文或者中文）
    # 2. 将母语内容填充在excel表里做为原文，然后通过调用机翻，将原文翻译为各语种的译文，并保存在excel表里
    # 3. 将各语种的译文作为模板的内容部分，“填充”到word模板里，并保存为id+langId的名称格式


    # 1. path是excel表的路径，sheetName是数据存放的tab名
    # 2. tempPath是word模板的路径，destFolder_path是生成的新的word文件保存的路径

    path=u"E:\\_测试组文件\\测试资源\\机翻测试.xlsx"
    sheetName = u'word文档翻译原文生成'
    tempPath = u"E:\\_00000\\TemplateWord.docx"
    destFolder_path = u"E:\\_00000\\TestDoc\\"




    #读取excel表，并获取原文，text1到text4（可自行扩展）
    workbook = ExcelUtils.ExcelDriver(path)
    worksheet = workbook.getWorksheet(sheetName)

    text1 = ExcelUtils.ExcelDriver.getCellData(2, 4, worksheet)
    text2 = ExcelUtils.ExcelDriver.getCellData(2, 5, worksheet)
    text3 = ExcelUtils.ExcelDriver.getCellData(2, 6, worksheet)
    text4 = ExcelUtils.ExcelDriver.getCellData(2, 7, worksheet)
    print text1


    #使用外部谷歌机翻进行翻译，并将译文填写到excel里，填写完成后保存Excel
    js=GoogleTranslateUtils.Py4Js()
    tk1 = js.getTk(text1)
    tk2 = js.getTk(text2)
    tk3 = js.getTk(text3)
    tk4 = js.getTk(text4)

    num = workbook.getSheetRowCount(worksheet)
    for index in range(num-1):
        langId = workbook.getCellData(index+3,3,worksheet)
        print langId
        transText1 = GoogleTranslateUtils.translate(text1, tk1, langId)
        transText2 = GoogleTranslateUtils.translate(text2, tk2, langId)
        transText3 = GoogleTranslateUtils.translate(text3, tk3, langId)
        transText4 = GoogleTranslateUtils.translate(text4, tk4, langId)
        print transText1
        resultCell1 = worksheet.cell(index + 3, 4)
        resultCell2 = worksheet.cell(index + 3, 5)
        resultCell3 = worksheet.cell(index + 3, 6)
        resultCell4 = worksheet.cell(index + 3, 7)

        workbook.setCellData(worksheet, resultCell1, transText1)
        workbook.setCellData(worksheet, resultCell2, transText2)
        workbook.setCellData(worksheet, resultCell3, transText3)
        workbook.setCellData(worksheet, resultCell4, transText4)

    workbook.saveExcelFile(path)
    #print TempCell


    #再通过遍历Excel，将刚刚保存的各语种的译文通过word模板，以context为规则填充到模板里，并保存到新的文件路径
    for index in range(num-1):
        langId = workbook.getCellData(index + 3, 3, worksheet)
        caseId = workbook.getCellData(index + 3, 1, worksheet)

        paraStr1 = workbook.getCellData(index + 3, 4, worksheet)
        paraStr2 = workbook.getCellData(index + 3, 5, worksheet)
        paraStr3 = workbook.getCellData(index + 3, 6, worksheet)
        paraStr4 = workbook.getCellData(index + 3, 7, worksheet)

        dest_path = destFolder_path+str(caseId)+'_'+langId+'.docx'
        context = {'paragraph1': paraStr1, 'list1': paraStr2,'list2': paraStr3, 'list3':paraStr4 }
        WordUtils.DocUtils.applyTempleDoc(tempPath, dest_path, context)
    #paraStr = workbook.getCellData(index + 3, 3, worksheet)
    #context = {'paragraph1': paraStr }
    # print os.path.exists(path)
    # print os.path.exists(dest_path)
    # pathlist = os.path.split(dest_path)
    # folderPath = pathlist[0]
    # fileName = pathlist[1]
    # print folderPath
    # print fileName

    # file_tmp = docx.Document(path)
    #
    # #file_tmp = docx.Document()
    # file_tmp.add_paragraph(u'第一段')
    # file_tmp.add_paragraph(u'第二段')
    #
    # file_tmp.save(dest_path)