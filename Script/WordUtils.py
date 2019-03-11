# -*- coding: utf-8 -*-
'''
Transn is pleased to support the open source community by making this file available.
Copyright (C) 2018 THL A29 Limited, a Tencent company. All rights reserved.
Licensed under the BSD 3-Clause License (the "License"); you may not use this file except in compliance with the License. You may obtain a copy of the License at
https://opensource.org/licenses/BSD-3-Clause
Unless required by applicable law or agreed to in writing, software distributed under the License is distributed on an "AS IS" BASIS, WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied. See the License for the specific language governing permissions and limitations under the License.

'''



import openpyxl
import docx
from docxtpl import DocxTemplate,RichText
import os
import re
import shutil
import lxml
from nt import chdir

class DocUtils:

    def __init__(self):
        pass
        #try:
            # self.workbook = xlrd.open_workbook(path)
            #file = docx.Document(path)

        #except:

             #raise RuntimeError()
    @staticmethod
    def showDocString(path):
        file = docx.Document(path)
        print("段落数:" + str(len(file.paragraphs)))
        for para in file.paragraphs:
            print(para.text)
        #DocUtils(u'E:\\_00000\\test01.docx')

    @staticmethod
    def createDocxFile(path):
        file = docx.Document()
        file.add_paragraph(u"啊")
        file.add_paragraph(u"Ah")
        file.save(path)


    @staticmethod
    def applyTempleDoc(path,path_new,context):
        doc = DocxTemplate(path)
        #context={'test_space' : '          ','test_tabs': 5*'\t','test_space_r' : RichText('          '),'test_tabs_r': RichText(5*'\t'),}
        doc.render(context)
        doc.save(path_new)


if __name__ == "__main__":
    #DocUtils.showDocString(u'E:\\_00000\\test01.docx')
    #DocUtils.createDocxFile(u'E:\\_00000\\test01_new.docx')
    context = {'company':u"222微软国际有限公司"}
    DocUtils.applyTempleDoc(u'E:\\_00000\\test01.docx',u'E:\\_00000\\test01_new.docx',context)
